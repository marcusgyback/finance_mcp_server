"""
tools/storage.py — 7 storage tools registered on the FastMCP instance.

All tools are async and interact with cloud storage (S3 / Filestash) through
the shared StorageClient abstraction.  Call register_tools(mcp) from server.py
to attach every tool to the FastMCP instance.
"""

from __future__ import annotations

import io
import json
import os
import tempfile
from typing import Optional

import openpyxl
import pdfplumber
from docx import Document
from mcp.server.fastmcp import FastMCP

from auth.context import client_folder
from storage.client import StorageClient


def _storage() -> StorageClient:
    """Return a fresh StorageClient (reads env vars on each call so the server
    can be restarted with new credentials without code changes)."""
    return StorageClient()


def _scoped_path(path: str) -> str:
    """Prepend the client's folder prefix to *path*.

    Also strips any leading slashes and collapses `../` traversal attempts so
    a client can never escape their own folder.

    Example:
        client folder = "acme-corp"
        path          = "reports/Q1.pdf"
        result        = "acme-corp/reports/Q1.pdf"
    """
    folder = client_folder.get()

    # Remove leading slashes and collapse path traversal sequences
    safe_path = path.lstrip("/")
    # Resolve traversal: split, drop empty and ".." parts, rejoin
    parts = [p for p in safe_path.replace("\\", "/").split("/") if p and p != ".."]
    safe_path = "/".join(parts)

    if folder:
        return f"{folder.rstrip('/')}/{safe_path}" if safe_path else folder
    return safe_path


def register_tools(mcp: FastMCP) -> None:
    """Attach all 7 storage tools to *mcp*."""

    # ------------------------------------------------------------------
    # Tool 1 — list_folder
    # ------------------------------------------------------------------
    @mcp.tool()
    async def list_folder(path: str = "") -> str:
        """List files and folders stored under the given path prefix.

        Returns a formatted directory listing with file sizes and last-modified
        timestamps.  Pass an empty string (the default) to list the bucket root.

        Args:
            path: Folder prefix to list, e.g. "reports/2024/".  Trailing slash
                  is optional.
        """
        client = _storage()
        scoped = _scoped_path(path)
        objects = client.list_objects(prefix=scoped)

        if not objects:
            return f"No files found under path: '{path}'"

        # Strip the client folder prefix from displayed keys so clients see
        # paths relative to their own root.
        folder = client_folder.get() or ""
        prefix_to_strip = f"{folder.rstrip('/')}/" if folder else ""
        for obj in objects:
            obj["display_key"] = obj["key"].removeprefix(prefix_to_strip)

        lines = [f"Contents of '{path}':", ""]
        col_key = max(len(o["display_key"]) for o in objects)
        col_size = max(len(str(o["size"])) for o in objects)
        header = f"{'Name':<{col_key}}  {'Size (bytes)':>{col_size}}  Last Modified"
        lines.append(header)
        lines.append("-" * len(header))

        for obj in objects:
            last_mod = obj["last_modified"].strftime("%Y-%m-%d %H:%M:%S UTC")
            lines.append(
                f"{obj['display_key']:<{col_key}}  {obj['size']:>{col_size}}  {last_mod}"
            )

        lines.append(f"\n{len(objects)} object(s) found.")
        return "\n".join(lines)

    # ------------------------------------------------------------------
    # Tool 2 — read_pdf
    # ------------------------------------------------------------------
    @mcp.tool()
    async def read_pdf(
        path: str,
        page_start: int = 1,
        page_end: Optional[int] = None,
    ) -> str:
        """Read a PDF file from storage and return its text content.

        For large PDFs, use page_start and page_end to read in chunks.
        Always check total_pages in the response and make multiple calls
        if needed to cover the full document.

        Args:
            path:       Storage key of the PDF file.
            page_start: First page to read (1-based, default: 1).
            page_end:   Last page to read inclusive. If omitted, reads up to
                        50 pages from page_start.
        """
        client = _storage()
        raw = client.get_object(_scoped_path(path))

        output_sections: list[str] = []

        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            total_pages = len(pdf.pages)
            start_idx = max(0, page_start - 1)
            end_idx = min(total_pages, (page_end or page_start + 49))
            pages_to_read = pdf.pages[start_idx:end_idx]

            output_sections.append(
                f"PDF: {path} | Total pages: {total_pages} | "
                f"Reading pages {start_idx + 1}–{end_idx}"
            )

            for page_num, page in enumerate(pages_to_read, start=start_idx + 1):
                output_sections.append(f"=== Page {page_num} ===")

                text = page.extract_text()
                if text and text.strip():
                    output_sections.append(text.strip())

                tables = page.extract_tables()
                for table_idx, table in enumerate(tables, start=1):
                    output_sections.append(f"\n[Table {table_idx} — Page {page_num}]")
                    output_sections.append(_format_table(table))

            if end_idx < total_pages:
                output_sections.append(
                    f"\n--- {total_pages - end_idx} more pages remaining. "
                    f"Call read_pdf with page_start={end_idx + 1} to continue. ---"
                )

        if not output_sections:
            return "No text content could be extracted from this PDF."

        return "\n\n".join(output_sections)

    # ------------------------------------------------------------------
    # Tool 3 — read_spreadsheet
    # ------------------------------------------------------------------
    @mcp.tool()
    async def read_spreadsheet(
        path: str,
        sheet_name: Optional[str] = None,
        row_start: int = 1,
        row_end: Optional[int] = None,
    ) -> str:
        """Read an Excel workbook (.xlsx) from storage and return its data as
        CSV-formatted text.

        For large spreadsheets, use row_start and row_end to read in chunks.
        Always check total_rows in the response and make multiple calls if
        needed to cover all data.

        Args:
            path:       Storage key of the .xlsx file.
            sheet_name: Name of a specific sheet to read. If omitted, reads
                        the first sheet.
            row_start:  First row to read (1-based, default: 1).
            row_end:    Last row to read inclusive. If omitted, reads up to
                        1000 rows from row_start.
        """
        client = _storage()
        raw = client.get_object(_scoped_path(path))

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            tmp.write(raw)
            tmp_path = tmp.name

        try:
            wb = openpyxl.load_workbook(tmp_path, data_only=True, read_only=True)
            name = sheet_name or wb.sheetnames[0]

            if name not in wb.sheetnames:
                return f"Sheet '{name}' not found. Available sheets: {', '.join(wb.sheetnames)}"

            ws = wb[name]
            total_rows = ws.max_row or 0
            end = min(total_rows, row_end or row_start + 999)

            sections: list[str] = [
                f"File: {path} | Sheet: {name} | Total rows: {total_rows} | "
                f"Reading rows {row_start}–{end}",
                f"=== Sheet: {name} ===",
            ]

            rows: list[str] = []
            for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                if row_idx < row_start:
                    continue
                if row_idx > end:
                    break
                cells = ["" if cell is None else str(cell) for cell in row]
                rows.append(",".join(_csv_escape(c) for c in cells))

            sections.append("\n".join(rows))

            if end < total_rows:
                sections.append(
                    f"\n--- {total_rows - end} more rows remaining. "
                    f"Call read_spreadsheet with row_start={end + 1} to continue. ---"
                )
        finally:
            os.unlink(tmp_path)

        return "\n\n".join(sections)

    # ------------------------------------------------------------------
    # Tool 4 — read_file
    # ------------------------------------------------------------------
    @mcp.tool()
    async def read_file(path: str) -> str:
        """Read a file from storage and return its content as a string.

        Supported formats:
        - .docx  — Word document, returns plain text with paragraph breaks
        - .json  — pretty-printed
        - .csv   — returned as-is
        - .txt   — returned as-is

        Args:
            path: Storage key of the file to read.
        """
        client = _storage()
        raw = client.get_object(_scoped_path(path))
        extension = path.rsplit(".", 1)[-1].lower() if "." in path else ""

        if extension == "docx":
            doc = Document(io.BytesIO(raw))
            lines = []
            for para in doc.paragraphs:
                lines.append(para.text)
            return "\n".join(lines)

        text = raw.decode("utf-8", errors="replace")

        if extension == "json":
            try:
                parsed = json.loads(text)
                return json.dumps(parsed, indent=2, ensure_ascii=False)
            except json.JSONDecodeError:
                return text

        # .csv and .txt (and anything else) — return verbatim
        return text

    # ------------------------------------------------------------------
    # Tool 5 — write_file
    # ------------------------------------------------------------------
    @mcp.tool()
    async def write_file(path: str, content: str) -> str:
        """Write a text file to storage.

        If the content is valid JSON it is automatically pretty-printed before
        saving.  The file is created or overwritten if it already exists.

        Args:
            path:    Storage key where the file should be saved, e.g.
                     "outputs/result.json".
            content: Text content to write.
        """
        client = _storage()

        # Auto-detect JSON and pretty-print
        text_to_save = content
        if path.rsplit(".", 1)[-1].lower() == "json" or _looks_like_json(content):
            try:
                parsed = json.loads(content)
                text_to_save = json.dumps(parsed, indent=2, ensure_ascii=False)
            except json.JSONDecodeError:
                pass  # Leave content unchanged if it fails to parse

        client.put_object(_scoped_path(path), text_to_save.encode("utf-8"))
        return f"File written successfully to: {path}"

    # ------------------------------------------------------------------
    # Tool 6 — create_spreadsheet
    # ------------------------------------------------------------------
    @mcp.tool()
    async def create_spreadsheet(
        path: str,
        data: str,
        sheet_name: str = "Sheet1",
    ) -> str:
        """Create an Excel (.xlsx) spreadsheet from JSON data and save it to
        storage.

        The *data* parameter must be a JSON string containing either:
          - A list of dicts  — each dict is a row; keys become column headers.
          - A list of lists  — each inner list is a row; no automatic headers.

        Args:
            path:       Destination storage key, e.g. "outputs/report.xlsx".
            data:       JSON-encoded table data (list of dicts or list of lists).
            sheet_name: Name of the worksheet (default: "Sheet1").
        """
        client = _storage()

        try:
            rows = json.loads(data)
        except json.JSONDecodeError as exc:
            return f"Error: data is not valid JSON — {exc}"

        if not isinstance(rows, list):
            return "Error: data must be a JSON array (list of dicts or list of lists)."

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name

        if rows and isinstance(rows[0], dict):
            # List of dicts — write headers from the first row's keys
            headers = list(rows[0].keys())
            ws.append(headers)
            for row_dict in rows:
                ws.append([row_dict.get(h) for h in headers])
        elif rows and isinstance(rows[0], list):
            for row_list in rows:
                ws.append(row_list)
        else:
            return "Error: data must be a non-empty list of dicts or list of lists."

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        client.put_object(_scoped_path(path), buffer.read())
        return f"Spreadsheet created successfully at: {path}"

    # ------------------------------------------------------------------
    # Tool 7 — create_word_document
    # ------------------------------------------------------------------
    @mcp.tool()
    async def create_word_document(path: str, content: str) -> str:
        """Create a Word document (.docx) from plain text with lightweight
        Markdown-style heading syntax and save it to storage.

        Heading syntax:
            # Title        → Heading 1
            ## Section     → Heading 2
            ### Subsection → Heading 3
            (blank line)   → paragraph break
            anything else  → normal paragraph

        Args:
            path:    Destination storage key, e.g. "outputs/report.docx".
            content: Document content using the heading syntax described above.
        """
        client = _storage()
        doc = Document()

        for line in content.splitlines():
            stripped = line.rstrip()

            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped == "":
                # Blank line — add an empty paragraph as a visual break
                doc.add_paragraph("")
            else:
                doc.add_paragraph(stripped)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        client.put_object(_scoped_path(path), buffer.read())
        return f"Word document created successfully at: {path}"


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------

def _format_table(table: list[list]) -> str:
    """Format a pdfplumber table as left-aligned columns.

    None cells are replaced with empty strings so the formatting never raises
    a TypeError.
    """
    if not table:
        return "(empty table)"

    # Normalise: replace None with ""
    str_table = [
        [("" if cell is None else str(cell)) for cell in row]
        for row in table
    ]

    # Calculate column widths
    num_cols = max(len(row) for row in str_table)
    col_widths = [0] * num_cols
    for row in str_table:
        for col_idx, cell in enumerate(row):
            if col_idx < num_cols:
                col_widths[col_idx] = max(col_widths[col_idx], len(cell))

    # Build formatted rows
    lines: list[str] = []
    for row in str_table:
        padded = []
        for col_idx in range(num_cols):
            cell = row[col_idx] if col_idx < len(row) else ""
            padded.append(cell.ljust(col_widths[col_idx]))
        lines.append("  ".join(padded).rstrip())

    return "\n".join(lines)


def _csv_escape(value: str) -> str:
    """Wrap a cell value in quotes if it contains a comma, newline, or quote."""
    if "," in value or '"' in value or "\n" in value:
        return '"' + value.replace('"', '""') + '"'
    return value


def _looks_like_json(text: str) -> bool:
    """Return True if the text appears to be JSON (starts with { or [)."""
    stripped = text.lstrip()
    return stripped.startswith("{") or stripped.startswith("[")
